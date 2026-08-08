from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from werkzeug.exceptions import HTTPException
from werkzeug.utils import secure_filename
from docx import Document
from playwright.sync_api import sync_playwright
from datetime import datetime
import os, json, uuid, html as html_module, base64, re, tempfile, zipfile, io
import xml.etree.ElementTree as ET
import openai
app = Flask(__name__)
CORS(app)
@app.errorhandler(HTTPException)
def handle_http_exception(e):
    if request.path.startswith('/api/'):
        return jsonify({"error": e.description}), e.code
    return e
@app.errorhandler(Exception)
def handle_unexpected_exception(e):
    if request.path.startswith('/api/'):
        app.logger.exception(e)
        return jsonify({"error": str(e)}), 500
    raise e
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
DATA_FOLDER = os.path.join(BASE_DIR, 'data')
EXPORT_FOLDER = os.path.join(BASE_DIR, 'exports')
FRONTEND_PATH = os.path.join(os.path.dirname(BASE_DIR), 'front')
for folder in [UPLOAD_FOLDER, DATA_FOLDER, EXPORT_FOLDER]:
    os.makedirs(folder, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
ALLOWED_EXTENSIONS = {'docx', 'doc'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
def get_data_path(doc_id):
    return os.path.join(DATA_FOLDER, f"{doc_id}.json")
def load_doc(doc_id):
    path = get_data_path(doc_id)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return json.JSONDecoder().raw_decode(text.lstrip())[0]
def save_doc(doc_id, data):
    data['updated_at'] = datetime.now().isoformat()
    path = get_data_path(doc_id)
    fd, tmp = tempfile.mkstemp(dir=DATA_FOLDER, suffix='.tmp')
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        os.replace(tmp, path)
    except BaseException:
        if os.path.exists(tmp):
            os.remove(tmp)
        raise
# Namespace prefixes that some non-Word tools (converters, older office suites) emit inside
# word/document.xml without declaring on the root element. lxml's parser treats an undeclared
# prefix as a fatal, unrecoverable XML error ("Namespace prefix X on Y is not defined"), so a
# file that Word itself opens fine can still fail here. Re-declaring any missing ones on the
# root <w:document> element before parsing repairs this without altering document content.
_OOXML_EXTRA_NAMESPACES = {
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
    'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
    'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
    'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
    'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
}
# def _repair_missing_namespaces(xml_bytes):
#     xml = xml_bytes.decode('utf-8')
#     root_match = re.search(r'<w:document\b[^>]*>', xml)
#     if not root_match:
#         return xml_bytes
#     root_tag = root_match.group(0)
#     missing = {
#         prefix: uri for prefix, uri in _OOXML_EXTRA_NAMESPACES.items()
#         if f'xmlns:{prefix}=' not in root_tag and re.search(rf'[<\s/]{prefix}:', xml)
#     }
#     if not missing:
#         return xml_bytes
#     injected = ''.join(f' xmlns:{p}="{u}"' for p, u in missing.items())
#     repaired_root = root_tag[:-1] + injected + '>'
#     xml = xml[:root_match.start()] + repaired_root + xml[root_match.end():]
#     return xml.encode('utf-8')
# def _repair_docx(filepath):
#     with zipfile.ZipFile(filepath) as zin:
#         if 'word/document.xml' not in zin.namelist():
#             return None
#         entries = {name: zin.read(name) for name in zin.namelist()}
#     original = entries['word/document.xml']
#     repaired = _repair_missing_namespaces(original)
#     if repaired == original:
#         return None
#     entries['word/document.xml'] = repaired
#     buffer = io.BytesIO()
#     with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zout:
#         for name, data in entries.items():
#             zout.writestr(name, data)
#     buffer.seek(0)
#     return buffer
def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
    except Exception as e:
        if 'is not defined' not in str(e):
            raise
        repaired = _repair_docx(filepath)
        if repaired is None:
            raise
        doc = Document(repaired)
    content = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip('\n')
        if text.strip():
            content.append({'id': f'p-{i}', 'type': 'paragraph', 'text': text})
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip('\n')
                if text.strip():
                    content.append({'id': f't-{ti}-{ri}-{ci}', 'type': 'table-cell', 'text': text})
    return content
_CSS_PROP = {
    'highlight': 'background-color', 'textcolor': 'color', 'dropcap': 'color',
    'border': 'border-color', 'circle': 'border-color',
    'underline': 'text-decoration-color', 'wavyunderline': 'text-decoration-color',
    'strikethrough': 'text-decoration-color', 'overline': 'text-decoration-color',
    'fontsize': 'font-size', 'letterspacing': 'letter-spacing'
}
def _export_paths(doc_id, doc, version=1):
    base = os.path.splitext(doc['original_filename'])[0]
    suffix = '' if version <= 1 else f' ({version})'
    return {fmt: (f"{base}_enriched{suffix}.{fmt}", os.path.join(EXPORT_FOLDER, f"{doc_id}_{base}_enriched{suffix}.{fmt}")) for fmt in ('png', 'html', 'pdf')}
def _next_export_version(doc_id, doc):
    version = 1
    while os.path.exists(_export_paths(doc_id, doc, version)['png'][1]):
        version += 1
    return version
_LOCAL_FONT_FILES = {
    'Nanum Barun Gothic': ('NanumBarunGothic.woff2', '400'),
    'NanumSquare': ('NanumSquare.woff2', '400'),
    'MaruBuri': ('MaruBuri-Regular.woff2', '400'),
    'ChosunGs': ('ChosunGs.woff2', '400'),
    'Inter': ('Inter-Variable.woff2', '100 900'),
}
def _local_font_faces():
    fonts_dir = os.path.join(FRONTEND_PATH, 'fonts')
    faces = []
    for family, (filename, weight) in _LOCAL_FONT_FILES.items():
        path = os.path.join(fonts_dir, filename)
        if not os.path.exists(path) or os.path.getsize(path) == 0:
            continue
        with open(path, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode()
        faces.append(f"@font-face{{font-family:'{family}';src:url(data:font/woff2;base64,{b64}) format('woff2');font-weight:{weight};font-style:normal}}")
    return ''.join(faces)
def _icon_html(s, decoration_styles=None):
    decoration_styles = decoration_styles or []
    deco_classes = ['styled-text'] + [d['type'] for d in decoration_styles] if decoration_styles else []
    deco_inline = [f'--deco-{d["type"]}-color:{d["color"]}' for d in decoration_styles]
    style_attr = f' style="{";".join(deco_inline)}"' if deco_inline else ''
    cls = ['inline-icon inline-icon-emoji' if s.get('isEmoji') else 'inline-icon'] + deco_classes
    title_attr = f' title="{html_module.escape(s.get("iconName") or "", quote=True)}"'
    if s.get('svgCode'):
        inner = s['svgCode']
    elif s.get('iconData'):
        inner = f'<img src="{html_module.escape(s["iconData"], quote=True)}" alt="">'
    else:
        return ''
    return f'<span class="{" ".join(cls)}"{title_attr}{style_attr}>{inner}</span>'

_DOC_CSS = r"""
:root {
    --color-white: #fff;
    --color-bg: #fafafa;
    --color-text: #1a1a1a;
    --color-text-secondary: #666;
    --color-text-muted: #999;
    --color-border: #e0e0e0;
    --color-border-dark: #1a1a1a;
    --color-primary: #1a1a1a;
    --color-primary-hover: #333;
    --space-2xl: 32px;
    --font-sans: 'Inter', 'Nanum Barun Gothic', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    --radius-lg: 8px;
}
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: var(--font-sans); font-size: 14px; line-height: 1.5; color: var(--color-text); background: var(--color-bg); -webkit-font-smoothing: antialiased; }
.document-container { max-width: 800px; margin: 24px auto; background: var(--color-white); border: 1px solid var(--color-border); border-radius: var(--radius-lg); box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
.document-content { padding: var(--space-2xl) 48px; font-size: 15px; line-height: 1.8; color: var(--color-text); }
.document-content p { margin-bottom: 1em; position: relative; white-space: pre-line; }
.document-content p:last-child { margin-bottom: 0; }
.styled-text { position: relative; display: inline; }
.styled-text.bold { font-weight: 700; }
.styled-text.italic { font-style: italic; }
.styled-text.underline { text-decoration: underline; text-decoration-thickness: 1.5px; text-underline-offset: 2px; z-index: 1; }
.styled-text.strikethrough { text-decoration: line-through; text-decoration-thickness: 1.5px; }
.styled-text.link { color: #0066cc; text-decoration: underline; text-decoration-color: #0066cc; text-decoration-thickness: 1.5px; text-underline-offset: 2px; }
.styled-text.highlight { padding: 1px 2px; border-radius: 2px; box-decoration-break: clone; -webkit-box-decoration-break: clone; }
.styled-text.border { border: 1.5px solid; border-radius: 0.2em; padding: 0.08em 0.15em; margin: 0 0.15em; }
.styled-text.circle { border: 1.5px solid; border-radius: 100px; padding: 0.15em 0.3em; margin: 0 0.15em; }
.styled-text.border .styled-text.highlight { padding: 0.08em 0.15em; margin: -0.08em -0.15em; border-radius: 0.2em; }
.styled-text.circle .styled-text.highlight { padding: 0.15em 0.1em; margin: -0.15em -0.3em; border-radius: 100px; }
.styled-text.serif { font-family: Georgia, 'Times New Roman', serif; }
.styled-text.arial { font-family: Arial, sans-serif; }
.styled-text.courier { font-family: 'Courier New', Courier, monospace; }
.styled-text.georgia { font-family: Georgia, 'Times New Roman', serif; }
.styled-text.helvetica { font-family: Helvetica, Arial, sans-serif; }
.styled-text.times { font-family: 'Times New Roman', Times, serif; }
.styled-text.trebuchet { font-family: 'Trebuchet MS', Helvetica, sans-serif; }
.styled-text.verdana { font-family: Verdana, Geneva, sans-serif; }
.styled-text.comicsans { font-family: 'Comic Sans MS', 'Comic Sans', cursive; }
.styled-text.cursivefont { font-family: cursive; }
.styled-text.bongothic { font-family: 'Noto Sans KR', sans-serif; }
.styled-text.nanumgothic { font-family: 'Nanum Gothic', sans-serif; }
.styled-text.bonmyeongjo { font-family: 'Noto Serif KR', serif; }
.styled-text.nanummyeongjo { font-family: 'Nanum Myeongjo', serif; }
.styled-text.nanumbarungothic { font-family: 'Nanum Barun Gothic', sans-serif; }
.styled-text.nanumsquare { font-family: 'NanumSquare', sans-serif; }
.styled-text.maruburi { font-family: 'MaruBuri', serif; }
.styled-text.gungseo { font-family: 'ChosunGs', serif; }
.styled-text.sansserif { font-family: 'Helvetica Neue', Arial, sans-serif; }
.styled-text.rounded { font-family: 'Nunito', 'Varela Round', sans-serif; }
.styled-text.mono { font-family: 'SF Mono', 'Consolas', 'Monaco', 'Courier New', monospace; }
.styled-text.smallcaps { font-variant: small-caps; letter-spacing: 0.05em; }
.styled-text.superscript { font-size: 0.7em; vertical-align: super; }
.styled-text.subscript { font-size: 0.7em; vertical-align: sub; }
.styled-text.overline { text-decoration: overline; text-decoration-thickness: 1.5px; }
.styled-text.wavyunderline { text-decoration: underline wavy; text-decoration-thickness: 1.5px; text-underline-offset: 1.5px; z-index: 1; }
.styled-text.dropcap { float: left; font-size: 3.2em; line-height: 0.8; padding-right: 8px; padding-top: 4px; font-weight: 700; }
.document-content p.callout-block { border: 2px solid; border-radius: 8px; padding: 12px 16px; margin-bottom: 1em; white-space: pre-wrap; }
.document-content p.callout-block.callout-merge-top { margin-top: -2px; border-top: none; border-top-left-radius: 0; border-top-right-radius: 0; }
.document-content p.callout-block.callout-merge-bottom { margin-bottom: -1px; border-bottom: none; border-bottom-left-radius: 0; border-bottom-right-radius: 0; }
.document-content p.quote-marks { text-align: center; padding: 48px 32px; font-style: italic; position: relative; }
.document-content p.quote-marks::before { content: '\201C'; position: absolute; top: 2px; left: 0; right: 0; font-size: 42px; line-height: 1; color: #c0c0c0; font-family: Georgia, 'Times New Roman', serif; text-align: center; }
.document-content p.quote-marks::after  { content: '\201D'; position: absolute; bottom: 2px; left: 0; right: 0; font-size: 42px; line-height: 1; color: #c0c0c0; font-family: Georgia, 'Times New Roman', serif; text-align: center; }
.document-content p.quote-line { border-left: 3px solid #c0c0c0; padding: 4px 0 4px 18px; margin-left: 8px; color: #555; }
.document-content p.list-bullet,
.document-content p.list-numbered { padding-left: 28px; position: relative; }
.document-content p.list-bullet::before { content: '\2022'; position: absolute; left: 10px; color: #666; }
.document-content p.list-numbered::before { content: attr(data-list-num) '.'; position: absolute; left: 6px; color: #666; font-size: 0.9em; }
.document-content p.code-white,
.document-content p.code-gray,
.document-content p.code-black {
    font-family: 'SF Mono', 'Consolas', 'Monaco', 'Courier New', monospace;
    font-size: 0.85em;
    line-height: 1.6;
    white-space: pre-wrap;
    word-break: break-word;
    padding: 18px 16px 14px;
    border-radius: 8px;
    margin-top: 12px;
}
.document-content p.code-white { background: #f6f8fa; color: #24292e; border: 1px solid #e1e4e8; }
.document-content p.code-gray  { background: #e5e5e7; color: #1f1f1f; border: 1px solid #d0d0d3; }
.document-content p.code-black { background: #1e1e1e; color: #d4d4d4; border: 1px solid #000000; }
.document-content p.code-white::before,
.document-content p.code-gray::before,
.document-content p.code-black::before {
    content: attr(data-code-lang);
    position: absolute;
    top: -10px;
    right: 12px;
    font-family: var(--font-sans);
    font-size: 10px;
    font-weight: 600;
    letter-spacing: 0.03em;
    text-transform: uppercase;
    padding: 2px 8px;
    border-radius: 100px;
}
.document-content p.code-white::before,
.document-content p.code-gray::before { background: #ffffff; color: #666666; border: 1px solid #d8d8d8; }
.document-content p.code-black::before { background: #2d2d2d; color: #aaaaaa; border: 1px solid #000000; }
.inline-icon {
    --icon-lift: -0.05em;
    position: relative; display: inline-block; height: 1em; width: auto; vertical-align: var(--icon-lift); margin: 0 2px;
}
.inline-icon svg, .inline-icon img { height: 1em; width: auto; display: block; }
.inline-icon-emoji { --icon-lift: -0.2em; height: auto; font-size: 1.15em; line-height: 1; vertical-align: var(--icon-lift); }
.inline-icon.underline, .inline-icon.strikethrough, .inline-icon.overline, .inline-icon.wavyunderline { text-decoration: none; }
.inline-icon.underline::after,
.inline-icon.strikethrough::after,
.inline-icon.overline::after,
.inline-icon.wavyunderline::after {
    content: ''; position: absolute; left: 0; right: 0; height: 1.5px; pointer-events: none;
}
.inline-icon.underline::after { bottom: calc(-1 * (var(--icon-lift) + 3.5px)); background: var(--deco-underline-color, currentColor); }
.inline-icon.strikethrough::after { top: calc(50% + var(--icon-lift) + 2.45px); transform: translateY(-50%); background: var(--deco-strikethrough-color, currentColor); }
.inline-icon.overline::after { top: calc(var(--icon-lift) - 1.3px); background: var(--deco-overline-color, currentColor); }
.inline-icon.wavyunderline::after {
    bottom: calc(-1 * (var(--icon-lift) + 3.5px)); height: 4px; background-color: var(--deco-wavyunderline-color, currentColor);
    -webkit-mask-repeat: repeat-x; mask-repeat: repeat-x;
    -webkit-mask-size: 8px 4px; mask-size: 8px 4px;
    -webkit-mask-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='8' height='4'%3E%3Cpath d='M0,2 Q2,0 4,2 Q6,4 8,2' stroke='black' stroke-width='1.4' fill='none'/%3E%3C/svg%3E");
    mask-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='8' height='4'%3E%3Cpath d='M0,2 Q2,0 4,2 Q6,4 8,2' stroke='black' stroke-width='1.4' fill='none'/%3E%3C/svg%3E");
}
.inline-icon-emoji.underline::after { bottom: calc(-1 * (var(--icon-lift) + 1px)); }
.inline-icon-emoji.strikethrough::after { top: calc(50% + var(--icon-lift) + 1.2px); }
.inline-icon-emoji.wavyunderline::after { bottom: calc(-1 * (var(--icon-lift) + 4.25px)); }
.divider-block { display: flex; align-items: center; margin: 16px 0; }
.div-line { flex: 1; border-top: 1px solid #d0d0d0; }
.div-sym-svg { flex-shrink: 0; display: block; }
.div-vline { width: 1px; height: 40px; background: #d0d0d0; }
.div-slash { width: 1px; height: 40px; background: #d0d0d0; transform: rotate(20deg); }
.divider-short .div-line,
.divider-bold .div-line,
.divider-dotted .div-line { flex: none; width: 28%; }
.divider-bold .div-line { border-top-width: 3px; border-top-color: #aaa; }
.divider-dotted .div-line { border-top-style: dotted; border-top-width: 2px; border-top-color: #d0d0d0; }
.divider-block[data-align="left"]   { justify-content: flex-start; }
.divider-block[data-align="center"] { justify-content: center; }
.divider-block[data-align="right"]  { justify-content: flex-end; }
"""
def _divider_html(d):
    line = '<div class="div-line"></div>'
    triangle_svg = '<svg class="div-sym-svg" viewBox="0 0 20 18" width="20" height="18"><polygon points="0,9 20,9 10,17" fill="none" stroke="#d0d0d0" stroke-width="1"/></svg>'
    diamond_svg = '<svg class="div-sym-svg" viewBox="0 0 20 18" width="20" height="18"><polygon points="0,9 10,3 20,9 10,15" fill="none" stroke="#d0d0d0" stroke-width="1"/></svg>'
    inner = {
        'full': line, 'short': line, 'bold': line, 'dotted': line,
        'triangle': line + triangle_svg + line,
        'diamond': line + diamond_svg + line,
        'cross': '<div class="div-slash"></div>',
        'vertical': '<div class="div-vline"></div>',
    }.get(d.get('dividerType'), line)
    align = f' data-align="{d["alignment"]}"' if d.get('alignment') else ''
    return f'<div class="divider-block divider-{d.get("dividerType")}"{align}>{inner}</div>'
def _wrap_styled(styles_list, inner_html):
    if not styles_list:
        return inner_html
    classes = ['styled-text'] + [s['type'] for s in styles_list]
    inline = [f'{_CSS_PROP[s["type"]]}:{s["color"]}' for s in styles_list if s['type'] in _CSS_PROP]
    style_attr = f' style="{";".join(inline)}"' if inline else ''
    return f'<span class="{" ".join(classes)}"{style_attr}>{inner_html}</span>'
def _border_span_style(b, seg_styles):
    fs = next((st for st in seg_styles if st['type'] == 'fontsize'), None)
    if fs:
        return f'border-color:{b["color"]};font-size:{fs["color"]}'
    return f'border-color:{b["color"]}'
def _style_covers_icon(st, pos):
    if st['startOffset'] == st['endOffset']:
        return st['startOffset'] == pos
    if st['startOffset'] == pos:
        return st.get('startIconInclusive') is not False
    if st['endOffset'] == pos:
        return st.get('endIconInclusive') is not False
    return st['startOffset'] < pos < st['endOffset']
def build_styled_html(content, styles):
    border_types = {'border', 'circle'}
    decoration_types = {'underline', 'strikethrough', 'overline', 'wavyunderline'}
    block_types = {'inlineicon', 'callout', 'quote', 'list', 'code'}
    styles_by_para = {}
    for s in styles:
        if s['type'] == 'divider':
            continue
        styles_by_para.setdefault(s['paraIndex'], []).append(s)
    dividers_by_para = {}
    for d in styles:
        if d['type'] == 'divider':
            dividers_by_para.setdefault(d['paraIndex'], []).append(d)
    numbered_para_indices = sorted(
        pi for pi in {s['paraIndex'] for s in styles if s['type'] == 'list'}
        if [s for s in styles_by_para.get(pi, []) if s['type'] == 'list'][-1].get('listStyle') == 'numbered'
    )
    numbered_list_values = {}
    running = 0
    for pi in numbered_para_indices:
        ls = [s for s in styles_by_para[pi] if s['type'] == 'list'][-1]
        running = 1 if ls.get('restart') else running + 1
        numbered_list_values[pi] = running
    paragraphs = []
    for i, para in enumerate(content):
        text = para['text']
        para_styles = styles_by_para.get(i, [])
        if not para_styles:
            paragraphs.append({'kind': 'p', 'classes': [], 'style': [], 'attrs': '', 'inner': html_module.escape(text), 'callout_group': None})
            paragraphs.extend({'kind': 'divider', 'html': _divider_html(d)} for d in reversed(dividers_by_para.get(i, [])))
            continue
        n = len(text)
        icon_styles = [s for s in para_styles if s['type'] == 'inlineicon']
        callout_styles = [s for s in para_styles if s['type'] == 'callout']
        quote_styles = [s for s in para_styles if s['type'] == 'quote']
        list_styles = [s for s in para_styles if s['type'] == 'list']
        code_styles = [s for s in para_styles if s['type'] == 'code']
        text_styles = [s for s in para_styles if s['type'] not in block_types]
        border_styles = [s for s in text_styles if s['type'] in border_types]
        non_border_styles = [s for s in text_styles if s['type'] not in border_types]
        offsets = {0, n}
        for s in text_styles:
            offsets.add(max(0, min(s['startOffset'], n)))
            offsets.add(max(0, min(s['endOffset'], n)))
        for s in icon_styles:
            offsets.add(max(0, min(s['startOffset'], n)))
        bounds = sorted(offsets)
        segments = []
        for j in range(len(bounds) - 1):
            start, end = bounds[j], bounds[j + 1]
            for s in icon_styles:
                if s['startOffset'] == start:
                    icon_active = [st for st in non_border_styles if _style_covers_icon(st, start)]
                    icon_active_borders = [st for st in border_styles if _style_covers_icon(st, start)]
                    segments.append({'kind': 'icon', 'style': s, 'styles': icon_active, 'borders': icon_active_borders})
            seg = text[start:end]
            if not seg:
                continue
            active = [s for s in non_border_styles if s['startOffset'] <= start and s['endOffset'] >= end]
            active_borders = [s for s in border_styles if s['startOffset'] <= start and s['endOffset'] >= end]
            segments.append({'kind': 'text', 'text': seg, 'styles': active, 'borders': active_borders})
        parts = ''
        open_borders = []
        for seg in segments:
            new_border_ids = {s.get('id', id(s)) for s in seg['borders']}
            close_from = next((j for j, ob in enumerate(open_borders) if ob.get('id', id(ob)) not in new_border_ids), None)
            if close_from is not None:
                to_reopen = [ob for ob in open_borders[close_from + 1:] if ob.get('id', id(ob)) in new_border_ids]
                for _ in range(len(open_borders) - close_from):
                    parts += '</span>'
                open_borders = open_borders[:close_from]
                for b in to_reopen:
                    parts += f'<span class="styled-text {b["type"]}" style="{_border_span_style(b, seg["styles"])}">'
                    open_borders.append(b)
            for b in seg['borders']:
                bid = b.get('id', id(b))
                if not any(ob.get('id', id(ob)) == bid for ob in open_borders):
                    parts += f'<span class="styled-text {b["type"]}" style="{_border_span_style(b, seg["styles"])}">'
                    open_borders.append(b)
            if seg['kind'] == 'icon':
                deco_styles = [st for st in seg['styles'] if st['type'] in decoration_types]
                parts += _wrap_styled(seg['styles'], _icon_html(seg['style'], deco_styles))
                continue
            parts += _wrap_styled(seg['styles'], html_module.escape(seg['text']))
        for _ in open_borders:
            parts += '</span>'
        for s in icon_styles:
            if s['startOffset'] >= n:
                icon_active = [st for st in non_border_styles if _style_covers_icon(st, n)]
                deco_styles = [st for st in icon_active if st['type'] in decoration_types]
                parts += _wrap_styled(icon_active, _icon_html(s, deco_styles))
        p_classes = []
        p_inline = []
        p_attrs = ''
        callout_group = None
        if callout_styles:
            cs = callout_styles[-1]
            p_classes.append('callout-block')
            p_inline.append(f'border-color:{cs["color"]}')
            if cs.get('bgColor'):
                p_inline.append(f'background-color:{cs["bgColor"]}')
            callout_group = cs.get('calloutGroup') or cs.get('id') or f'__solo_{i}'
        if quote_styles:
            p_classes.append(f'quote-{quote_styles[-1]["quoteStyle"]}')
        if list_styles:
            ls = list_styles[-1]
            p_classes.append(f'list-{ls["listStyle"]}')
            if ls.get('listStyle') == 'numbered':
                p_attrs += f' data-list-num="{numbered_list_values.get(i, "")}"'
        if code_styles:
            cds = code_styles[-1]
            p_classes.append(f'code-{cds["bgStyle"]}')
            p_attrs += f' data-code-lang="{html_module.escape(cds.get("language", ""), quote=True)}"'
        paragraphs.append({'kind': 'p', 'classes': p_classes, 'style': p_inline, 'attrs': p_attrs, 'inner': parts, 'callout_group': callout_group})
        paragraphs.extend({'kind': 'divider', 'html': _divider_html(d)} for d in reversed(dividers_by_para.get(i, [])))

    for idx in range(len(paragraphs) - 1):
        cur, nxt = paragraphs[idx], paragraphs[idx + 1]
        if cur['kind'] != 'p' or nxt['kind'] != 'p':
            continue
        if 'callout-block' not in cur['classes'] or 'callout-block' not in nxt['classes']:
            continue
        if cur['callout_group'] is not None and cur['callout_group'] == nxt['callout_group']:
            cur['classes'].append('callout-merge-bottom')
            nxt['classes'].append('callout-merge-top')

    def _render_paragraph(item):
        if item['kind'] == 'divider':
            return item['html']
        class_attr = f' class="{" ".join(item["classes"])}"' if item['classes'] else ''
        style_attr = f' style="{";".join(item["style"])}"' if item['style'] else ''
        return f'<p{class_attr}{style_attr}{item["attrs"]}>{item["inner"]}</p>'

    body_html = ''.join(_render_paragraph(item) for item in paragraphs)
    return f'''<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Nunito:wght@400;500;600;700&family=Noto+Sans+KR:wght@400;500;700&family=Noto+Serif+KR:wght@400;500;700&family=Nanum+Gothic:wght@400;700&family=Nanum+Myeongjo:wght@400;700&display=swap" rel="stylesheet">
<style>
{_local_font_faces()}
{_DOC_CSS}
</style>
</head>
<body>
<div class="document-container"><div class="document-content">{body_html}</div></div>
</body>
</html>'''
@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "message": "Document Typography API is running"})
def _paragraphs_from_text(text):
    content = []
    for i, line in enumerate(text.split('\n')):
        line = line.strip('\n')
        if line.strip():
            content.append({'id': f'p-{i}', 'type': 'paragraph', 'text': line})
    return content
def _new_doc_data(doc_id, original_filename, content):
    return {
        "doc_id": doc_id, "original_filename": original_filename,
        "created_at": datetime.now().isoformat(), "updated_at": datetime.now().isoformat(),
        "content": content, "styles": [], "enrichment_log": []
    }
@app.route('/api/upload', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    if not allowed_file(file.filename):
        return jsonify({"error": "Only Word documents (.docx) are allowed"}), 400
    doc_id = str(uuid.uuid4())
    original_filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, f"{doc_id}.docx")
    file.save(file_path)
    try:
        content = extract_text_from_docx(file_path)
    except Exception as e:
        os.remove(file_path)
        return jsonify({"error": f"Failed to parse document: {str(e)}"}), 400
    doc_data = _new_doc_data(doc_id, original_filename, content)
    save_doc(doc_id, doc_data)
    return jsonify({"success": True, "doc_id": doc_id, "filename": original_filename, "content": content, "message": "Document uploaded successfully"})
@app.route('/api/upload-text', methods=['POST'])
def upload_text():
    data = request.get_json(silent=True) or {}
    text = data.get('text', '')
    if not text.strip():
        return jsonify({"error": "No text provided"}), 400
    original_filename = secure_filename((data.get('filename') or '').strip()) or 'Pasted Text'
    doc_id = str(uuid.uuid4())
    content = _paragraphs_from_text(text)
    doc_data = _new_doc_data(doc_id, original_filename, content)
    save_doc(doc_id, doc_data)
    return jsonify({"success": True, "doc_id": doc_id, "filename": original_filename, "content": content, "message": "Document created successfully"})
@app.route('/api/document/<doc_id>', methods=['GET'])
def get_document(doc_id):
    doc = load_doc(doc_id)
    return jsonify(doc) if doc else (jsonify({"error": "Document not found"}), 404)
@app.route('/api/document/<doc_id>/styles', methods=['POST'])
def save_styles(doc_id):
    doc = load_doc(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    data = request.get_json()
    doc['styles'] = data.get('styles', [])
    if data.get('content') is not None:
        doc['content'] = data['content']
    save_doc(doc_id, doc)
    return jsonify({"success": True, "message": "Styles saved successfully", "updated_at": doc['updated_at']})
@app.route('/api/document/<doc_id>/log', methods=['POST'])
def log_enrichment_action(doc_id):
    doc = load_doc(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    entry = request.get_json()
    if not entry or 'action' not in entry or 'timestamp' not in entry:
        return jsonify({"error": "Invalid log entry"}), 400
    if entry['action'] not in ('add', 'delete', 'clear'):
        return jsonify({"error": "Invalid action type"}), 400
    entry['id'] = doc.get('original_filename')
    doc.setdefault('enrichment_log', []).append(entry)
    save_doc(doc_id, doc)
    return jsonify({"success": True, "message": "Action logged", "log_count": len(doc['enrichment_log'])})

SVG_DISALLOWED_TAGS = {'script', 'foreignobject', 'iframe', 'image', 'style', 'animate', 'animatetransform', 'animatemotion', 'set'}

def validate_svg(svg_text):
    try:
        root = ET.fromstring(svg_text)
    except ET.ParseError:
        return False
    tag = root.tag.rsplit('}', 1)[-1].lower()
    if tag != 'svg':
        return False
    for el in root.iter():
        local = el.tag.rsplit('}', 1)[-1].lower()
        if local in SVG_DISALLOWED_TAGS:
            return False
        for attr, value in el.attrib.items():
            attr_local = attr.rsplit('}', 1)[-1].lower()
            if attr_local.startswith('on'):
                return False
            if attr_local in ('href', 'xlink:href') and not value.startswith('#'):
                return False
    return True

@app.route('/api/generate-icon', methods=['POST'])
def generate_icon():
    data = request.get_json()
    if not data or not data.get('description'):
        return jsonify({"error": "Description is required"}), 400
    description = data['description'].strip()
    if len(description) > 200:
        return jsonify({"error": "Description too long (max 200 characters)"}), 400
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    client = openai.OpenAI(api_key=api_key)
    system_prompt = (
        "You are an expert SVG icon designer. Create a single, high-quality SVG icon.\n"
        "Requirements:\n"
        "- Output ONLY valid, complete, well-formed, and intent-aligned SVG code. No explanation, no markdown, no comments.\n"
        "- Use <svg xmlns=\"http://www.w3.org/2000/svg\" width=\"24\" height=\"24\" viewBox=\"0 0 24 24\" "
        "stroke-width=\"2\" stroke-linecap=\"round\" stroke-linejoin=\"round\">\n"
        "- Color rules:\n"
        "- If the user mentions a color (e.g., 'red star', 'blue arrow'), use that exact color as hex values.\n "
        "- If no color is mentioned, use stroke=\"currentColor\" fill=\"none\" as defaults on the <svg> element.\n"
        "- The icon must be immediately recognizable when displayed at 16px size.\n"
        "- Keep the design clean with minimal elements. Avoid excessive detail.\n"
        "- Stay within coordinates 2-22 to ensure padding inside the viewBox.\n"
        "- Use well-known visual metaphors (e.g., checkmark=success, star=favorite, heart=love, gear=settings).\n"
        "- Every path/shape must be fully closed and syntactically complete; double-check before answering.\n"
    )

    def request_svg(extra_note=None):
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Generate an SVG icon for: {description}"},
        ]
        if extra_note:
            messages.append({"role": "user", "content": extra_note})
        response = client.chat.completions.create(
            model="gpt-5.2",
            messages=messages,
            max_completion_tokens=800,
            temperature=0.5
        )
        text = response.choices[0].message.content.strip()
        if text.startswith('```'):
            text = re.sub(r'^```(?:svg|xml)?\s*\n?', '', text)
            text = re.sub(r'\n?```\s*$', '', text)
            text = text.strip()
        return text

    try:
        svg_text = request_svg()
        if not validate_svg(svg_text):
            svg_text = request_svg("Your previous response was not valid, complete SVG. Return only one single well-formed <svg>...</svg> element.")
        if not validate_svg(svg_text):
            return jsonify({"error": "Failed to generate valid SVG"}), 500
        data_url = f"data:image/svg+xml;base64,{base64.b64encode(svg_text.encode()).decode()}"
        return jsonify({
            "success": True,
            "iconData": data_url,
            "iconName": description[:50],
            "svgCode": svg_text
        })
    except openai.APIError as e:
        return jsonify({"error": f"OpenAI API error: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Failed to generate icon: {str(e)}"}), 500
@app.route('/api/document/<doc_id>/export', methods=['POST'])
def export_document(doc_id):
    doc = load_doc(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    data = request.get_json()
    styles = data.get('styles', [])
    doc['styles'] = styles
    if data.get('content') is not None:
        doc['content'] = data['content']
    save_doc(doc_id, doc)
    export_paths = _export_paths(doc_id, doc, _next_export_version(doc_id, doc))
    export_filename, export_path = export_paths['png']
    html_filename, html_path = export_paths['html']
    pdf_filename, pdf_path = export_paths['pdf']
    try:
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(build_styled_html(doc['content'], styles))
        with sync_playwright() as p:
            browser = p.chromium.launch(args=['--no-sandbox', '--disable-dev-shm-usage', '--disable-gpu'] if os.environ.get('FLASK_ENV') == 'production' else [])
            page = browser.new_page(viewport={'width': 900, 'height': 800}, device_scale_factor=2)
            page.goto(f'file://{html_path}')
            page.wait_for_load_state('networkidle')
            page.wait_for_timeout(500)
            page.screenshot(path=export_path, full_page=True)
            page.emulate_media(media="screen")
            page.pdf(path=pdf_path, print_background=True, format='A4')
            browser.close()
        doc.setdefault('enrichment_log', []).append({
            'id': doc.get('original_filename'),
            'action': 'export',
            'timestamp': datetime.now().isoformat(),
            'files': {'png': export_filename, 'html': html_filename, 'pdf': pdf_filename}
        })
        save_doc(doc_id, doc)
    except Exception as e:
        return jsonify({"success": False, "error": f"Export failed: {str(e)}"}), 500
    return jsonify({"success": True, "message": "Document exported successfully", "export_filename": export_filename, "download_url": f"/api/document/{doc_id}/download"})
@app.route('/api/document/<doc_id>/download', methods=['GET'])
def download_document(doc_id):
    doc = load_doc(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    latest_version = _next_export_version(doc_id, doc) - 1
    if latest_version < 1:
        return jsonify({"error": "Export not found. Please save the document first."}), 404
    export_filename, export_path = _export_paths(doc_id, doc, latest_version)['png']
    if not os.path.exists(export_path):
        return jsonify({"error": "Export not found. Please save the document first."}), 404
    return send_file(export_path, as_attachment=True, download_name=export_filename)
@app.route('/api/documents', methods=['GET'])
def list_documents():
    docs = []
    for f in os.listdir(DATA_FOLDER):
        if f.endswith('.json'):
            with open(os.path.join(DATA_FOLDER, f), 'r', encoding='utf-8') as file:
                d = json.load(file)
                docs.append({"doc_id": d['doc_id'], "filename": d['original_filename'], "created_at": d['created_at'], "updated_at": d['updated_at']})
    return jsonify(docs)
@app.route('/api/document/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    for path in [os.path.join(UPLOAD_FOLDER, f"{doc_id}.docx"), get_data_path(doc_id)]:
        if os.path.exists(path):
            os.remove(path)
    return jsonify({"success": True, "message": "Document deleted successfully"})
@app.route('/')
def serve_frontend():
    return send_from_directory(FRONTEND_PATH, 'index.html')
@app.route('/<path:filename>')
def serve_static(filename):
    return send_from_directory(FRONTEND_PATH, filename)
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5002))
    print(f"{'=' * 50}\nServer running at: http://localhost:{port}\n{'=' * 50}")
    app.run(debug=os.environ.get('FLASK_ENV') != 'production', host='0.0.0.0', port=port)
